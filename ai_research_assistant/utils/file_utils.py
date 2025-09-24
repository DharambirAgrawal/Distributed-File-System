import math
import os
from io import BytesIO
import mimetypes

import docx
from PyPDF2 import PdfReader
from PIL import Image
import pytesseract

def get_user_upload_dir(user_id):
    """Get the upload directory for a specific user"""
    base_dir = "/workspaces/Distributed-File-System/ai_research_assistant/uploads"
    return os.path.join(base_dir, f"user_{user_id}")

def ensure_user_directory(directory):
    """Ensure that a directory exists, create if it doesn't"""
    if not os.path.exists(directory):
        os.makedirs(directory, exist_ok=True)

def get_file_size(file_path):
    """Get file size in bytes"""
    try:
        return os.path.getsize(file_path)
    except OSError:
        return 0

def format_file_size(size_bytes):
    """Format file size in human readable format"""
    if size_bytes == 0:
        return "0 B"
    
    size_names = ["B", "KB", "MB", "GB", "TB"]
    i = 0
    while size_bytes >= 1024.0 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1
    
    return f"{size_bytes:.1f} {size_names[i]}"

def extract_text_from_file(file_path):
    """Extract text content from various file types"""
    try:
        # Get file extension
        _, ext = os.path.splitext(file_path.lower())
        
        if ext == '.txt':
            return extract_text_from_txt(file_path)
        elif ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif ext in ['.doc', '.docx']:
            return extract_text_from_docx(file_path)
        elif ext in ['.png', '.jpg', '.jpeg', '.gif']:
            return extract_text_from_image(file_path)
        else:
            return "Unsupported file type for text extraction"
            
    except Exception as e:
        return f"Error extracting text: {str(e)}"

def extract_text_from_txt(file_path):
    """Extract text from TXT files"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except UnicodeDecodeError:
        # Try with different encoding
        with open(file_path, 'r', encoding='latin-1') as file:
            return file.read()

def extract_text_from_pdf(file_path):
    """Extract text from PDF files"""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_text_from_docx(file_path):
    """Extract text from DOCX files"""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

def extract_text_from_image(file_path):
    """Extract text from images using OCR"""
    try:
        # Use OCR to extract text from images
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        return f"OCR not available or error: {str(e)}"


def get_page_count_from_bytes(file_bytes: bytes, extension: str) -> int:
    """Estimate page count from in-memory file content."""
    ext = extension.lower().lstrip(".")

    if ext == "pdf":
        reader = PdfReader(BytesIO(file_bytes))
        return len(reader.pages)

    if ext == "docx":
        document = docx.Document(BytesIO(file_bytes))
        word_count = sum(len(paragraph.text.split()) for paragraph in document.paragraphs)
        table_word_count = 0
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_word_count += len(cell.text.split())
        total_words = word_count + table_word_count
        if total_words == 0:
            return 1
        return max(1, math.ceil(total_words / 300))

    if ext == "doc":
        raise ValueError("Page counting for legacy .doc files isn't supported. Please convert the document to PDF or DOCX.")

    raise ValueError(f"Page count checks are not supported for .{ext} files.")